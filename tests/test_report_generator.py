"""Tests for the DOCX report generator."""

import os

import pytest
from docx import Document

from models.reports import AnalysisSummary, VisualizationResult
from models.threats import ClassifiedThreat, ThreatReport
from utils.report_generator import ReportGenerator


@pytest.fixture
def threat_report_with_threats():
    threats = [
        ClassifiedThreat(
            category="PORT_SCAN",
            severity_score=85,
            severity_label="critical",
            title="Vertical Port Scan from 192.168.1.100",
            description="Vertical port scan detected targeting 25 ports",
            detection_method="rule",
            source_ips=["192.168.1.100"],
            dest_ips=["10.0.0.1"],
            time_range="2026-03-15T10:00:00 \u2013 2026-03-15T10:01:00",
            evidence={"scan_type": "vertical", "ports_scanned": 25},
            recommendations=["Block source IP at firewall"],
        ),
        ClassifiedThreat(
            category="SUSPICIOUS_CONNECTION",
            severity_score=45,
            severity_label="medium",
            title="Suspicious Connection to 10.0.0.1",
            description="Traffic on suspicious port 4444",
            detection_method="rule",
            source_ips=["192.168.1.50"],
            dest_ips=["10.0.0.1"],
            time_range="2026-03-15T10:00:00 \u2013 2026-03-15T10:00:00",
            evidence={"port": 4444},
            recommendations=["Isolate affected host for investigation"],
        ),
    ]
    return ThreatReport(
        total_threats=2,
        critical_count=1,
        high_count=0,
        medium_count=1,
        low_count=0,
        threats=threats,
        analysis_duration_seconds=5.0,
        packets_analyzed=1000,
        time_range_analyzed="2026-03-15T10:00:00 \u2013 2026-03-15T10:01:00",
    )


@pytest.fixture
def empty_threat_report():
    return ThreatReport(
        total_threats=0,
        critical_count=0,
        high_count=0,
        medium_count=0,
        low_count=0,
        threats=[],
        analysis_duration_seconds=1.0,
        packets_analyzed=100,
        time_range_analyzed="2026-03-15T10:00:00 \u2013 2026-03-15T10:00:10",
    )


@pytest.fixture
def analysis_summary(threat_report_with_threats):
    return AnalysisSummary(
        source_files=["test_capture.pcap"],
        total_packets=1000,
        total_log_entries=50,
        time_range="2026-03-15T10:00:00 \u2013 2026-03-15T10:01:00",
        unique_ips=15,
        protocol_breakdown={"TCP": 600, "UDP": 250, "ICMP": 100, "DNS": 50},
        threat_summary=threat_report_with_threats,
    )


@pytest.fixture
def empty_analysis_summary(empty_threat_report):
    return AnalysisSummary(
        source_files=["empty.pcap"],
        total_packets=100,
        total_log_entries=0,
        time_range="2026-03-15T10:00:00 \u2013 2026-03-15T10:00:10",
        unique_ips=5,
        protocol_breakdown={"TCP": 80, "UDP": 20},
        threat_summary=empty_threat_report,
    )


@pytest.fixture
def viz_result(tmp_path):
    return VisualizationResult(chart_dir=str(tmp_path), charts=[])


@pytest.fixture
def report_generator(tmp_path):
    output_dir = str(tmp_path / "reports")
    return ReportGenerator(output_dir=output_dir)


def test_generate_creates_docx(report_generator, analysis_summary, viz_result):
    filepath = report_generator.generate(analysis_summary, viz_result)
    assert os.path.exists(filepath)
    assert filepath.endswith(".docx")


def test_generate_file_not_empty(report_generator, analysis_summary, viz_result):
    filepath = report_generator.generate(analysis_summary, viz_result)
    assert os.path.getsize(filepath) > 0


def test_generate_filename_format(report_generator, analysis_summary, viz_result):
    filepath = report_generator.generate(analysis_summary, viz_result)
    filename = os.path.basename(filepath)
    assert filename.startswith("threat_analysis_")


def test_generate_with_zero_threats(report_generator, empty_analysis_summary, viz_result):
    filepath = report_generator.generate(empty_analysis_summary, viz_result)
    assert os.path.exists(filepath)
    assert os.path.getsize(filepath) > 0


def test_generate_creates_output_dir(tmp_path, analysis_summary, viz_result):
    nested_dir = str(tmp_path / "deep" / "nested" / "output")
    generator = ReportGenerator(output_dir=nested_dir)
    filepath = generator.generate(analysis_summary, viz_result)
    assert os.path.exists(filepath)


def test_report_contains_title(report_generator, analysis_summary, viz_result):
    filepath = report_generator.generate(analysis_summary, viz_result)
    doc = Document(filepath)
    text = "\n".join([p.text for p in doc.paragraphs])
    assert "NETWORK THREAT ANALYSIS REPORT" in text


def test_report_contains_methodology(report_generator, analysis_summary, viz_result):
    filepath = report_generator.generate(analysis_summary, viz_result)
    doc = Document(filepath)
    text = "\n".join([p.text for p in doc.paragraphs])
    assert "dual-layer detection approach" in text


def test_report_contains_threats(report_generator, analysis_summary, viz_result):
    filepath = report_generator.generate(analysis_summary, viz_result)
    doc = Document(filepath)
    text = "\n".join([p.text for p in doc.paragraphs])
    assert "Vertical Port Scan" in text
    assert "Suspicious Connection" in text
