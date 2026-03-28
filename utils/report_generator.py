"""DOCX threat report generator with embedded charts and formatted tables."""

import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from models.threats import ThreatReport, ClassifiedThreat
from models.reports import AnalysisSummary, VisualizationResult

DARK_BLUE = RGBColor(0x1A, 0x23, 0x7E)
ACCENT_RED = RGBColor(0xC6, 0x28, 0x28)
ACCENT_ORANGE = RGBColor(0xE6, 0x51, 0x00)
ACCENT_YELLOW = RGBColor(0xF9, 0xA8, 0x25)
ACCENT_TEAL = RGBColor(0x0D, 0x73, 0x77)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)

SEVERITY_COLORS = {
    "critical": "C62828",
    "high": "E65100",
    "medium": "F9A825",
    "low": "0D7377",
}


class ReportGenerator:

    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate(self, summary: AnalysisSummary, viz: VisualizationResult) -> str:
        doc = Document()
        self._apply_styles(doc)

        self._add_cover_page(doc, summary)
        self._add_executive_summary(doc, summary)
        self._add_traffic_overview(doc, summary)
        self._add_protocol_analysis(doc, summary, viz)
        self._add_threat_summary(doc, summary)
        self._add_threat_details(doc, summary)
        self._add_timeline_analysis(doc, viz)
        self._add_recommendations(doc, summary)
        self._add_methodology(doc)

        filename = f"threat_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def _apply_styles(self, doc: Document):
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        for level in range(1, 4):
            heading_style = doc.styles[f"Heading {level}"]
            heading_style.font.color.rgb = DARK_BLUE
            heading_style.font.name = "Calibri"

    def _add_cover_page(self, doc: Document, summary: AnalysisSummary):
        doc.add_paragraph("")
        doc.add_paragraph("")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("NETWORK THREAT ANALYSIS REPORT")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = DARK_BLUE

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(datetime.now().strftime("%B %d, %Y"))
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

        doc.add_paragraph("")

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = [
            f"Source files: {', '.join(summary.source_files)}",
            f"Total packets analyzed: {summary.total_packets:,}",
            f"Threats detected: {summary.threat_summary.total_threats}",
        ]
        for i, line in enumerate(lines):
            run = info.add_run(line)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)
            if i < len(lines) - 1:
                info.add_run("\n")

        doc.add_page_break()

    def _add_executive_summary(self, doc: Document, summary: AnalysisSummary):
        doc.add_heading("Executive Summary", level=1)

        ts = summary.threat_summary
        categories = {}
        for threat in ts.threats:
            categories[threat.category] = categories.get(threat.category, 0) + 1
        top_category = max(categories, key=categories.get) if categories else "N/A"

        text = (
            f"Analysis of {summary.total_packets:,} packets over {summary.time_range} "
            f"identified {ts.total_threats} threats "
            f"({ts.critical_count} critical, {ts.high_count} high). "
            f"{top_category} was the most prevalent threat category."
        )
        doc.add_paragraph(text)

    def _add_traffic_overview(self, doc: Document, summary: AnalysisSummary):
        doc.add_heading("Traffic Overview", level=1)

        table = doc.add_table(rows=4, cols=2)
        table.style = "Light Shading Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        rows_data = [
            ("Total Packets", f"{summary.total_packets:,}"),
            ("Unique IPs", f"{summary.unique_ips:,}"),
            ("Time Range", summary.time_range),
            ("Protocols", ", ".join(summary.protocol_breakdown.keys())),
        ]
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    def _add_protocol_analysis(
        self, doc: Document, summary: AnalysisSummary, viz: VisualizationResult,
    ):
        doc.add_heading("Protocol Analysis", level=1)

        total = sum(summary.protocol_breakdown.values()) or 1
        lines = []
        for proto, count in sorted(
            summary.protocol_breakdown.items(), key=lambda x: x[1], reverse=True,
        ):
            pct = count / total * 100
            lines.append(f"{proto}: {count:,} packets ({pct:.1f}%)")
        doc.add_paragraph("\n".join(lines))

        chart_path = os.path.join(viz.chart_dir, "protocol_distribution.png")
        if os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_threat_summary(self, doc: Document, summary: AnalysisSummary):
        doc.add_heading("Threat Summary", level=1)

        threats = summary.threat_summary.threats
        if not threats:
            doc.add_paragraph("No threats detected.")
            return

        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Severity", "Category", "Title", "Source IPs", "Detection"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.size = Pt(9)
            self._set_cell_bg(cell, "1A237E")

        for threat in threats:
            row = table.add_row()
            row.cells[0].text = threat.severity_label.upper()
            row.cells[1].text = threat.category
            row.cells[2].text = threat.title
            row.cells[3].text = ", ".join(threat.source_ips[:3])
            if len(threat.source_ips) > 3:
                row.cells[3].text += f" (+{len(threat.source_ips) - 3})"
            row.cells[4].text = threat.detection_method

            color = SEVERITY_COLORS.get(threat.severity_label, "757575")
            self._set_cell_bg(row.cells[0], color)

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    def _add_threat_details(self, doc: Document, summary: AnalysisSummary):
        doc.add_heading("Threat Details", level=1)

        threats = summary.threat_summary.threats[:10]
        for i, threat in enumerate(threats, 1):
            doc.add_heading(f"{i}. {threat.title}", level=2)
            doc.add_paragraph(threat.description)

            if threat.evidence:
                doc.add_heading("Evidence", level=3)
                ev_table = doc.add_table(rows=1, cols=2)
                ev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                ev_table.rows[0].cells[0].text = "Indicator"
                ev_table.rows[0].cells[1].text = "Value"
                for paragraph in ev_table.rows[0].cells[0].paragraphs:
                    paragraph.runs[0].bold = True
                for paragraph in ev_table.rows[0].cells[1].paragraphs:
                    paragraph.runs[0].bold = True

                for key, value in threat.evidence.items():
                    row = ev_table.add_row()
                    row.cells[0].text = str(key)
                    row.cells[1].text = str(value)

            if threat.source_ips:
                doc.add_heading("Affected IPs", level=3)
                doc.add_paragraph(", ".join(threat.source_ips))

            if threat.recommendations:
                doc.add_heading("Recommendations", level=3)
                for rec in threat.recommendations:
                    doc.add_paragraph(rec, style="List Bullet")

    def _add_timeline_analysis(self, doc: Document, viz: VisualizationResult):
        doc.add_heading("Timeline Analysis", level=1)

        chart_path = os.path.join(viz.chart_dir, "traffic_timeline.png")
        if os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        anomaly_path = os.path.join(viz.chart_dir, "anomaly_scores.png")
        if os.path.exists(anomaly_path):
            doc.add_picture(anomaly_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_recommendations(self, doc: Document, summary: AnalysisSummary):
        doc.add_heading("Recommendations", level=1)

        seen = set()
        prioritized = []
        severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3}

        sorted_threats = sorted(
            summary.threat_summary.threats,
            key=lambda t: severity_order.get(t.severity_label, 4),
        )
        for threat in sorted_threats:
            for rec in threat.recommendations:
                if rec not in seen:
                    seen.add(rec)
                    prioritized.append((threat.severity_label, rec))

        if not prioritized:
            doc.add_paragraph("No specific recommendations at this time.")
            return

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "#"
        table.rows[0].cells[1].text = "Priority"
        table.rows[0].cells[2].text = "Recommendation"
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].bold = True

        for i, (severity, rec) in enumerate(prioritized, 1):
            row = table.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = severity.upper()
            row.cells[2].text = rec
            color = SEVERITY_COLORS.get(severity, "757575")
            self._set_cell_bg(row.cells[1], color)

    def _add_methodology(self, doc: Document):
        doc.add_heading("Methodology", level=1)

        doc.add_paragraph(
            "This analysis employed a dual-layer detection approach combining "
            "rule-based signature matching with machine learning ensemble models."
        )
        doc.add_paragraph(
            "Rule-based detection applied protocol-aware pattern matching against "
            "known attack signatures, including port scan heuristics, volumetric "
            "thresholds for DDoS identification, and brute-force attempt correlation."
        )
        doc.add_paragraph(
            "The ML ensemble combined Isolation Forest, Local Outlier Factor, and "
            "One-Class SVM models trained on baseline traffic profiles. Anomaly "
            "scores were computed as weighted consensus across model votes, with "
            "a detection threshold of 0.5."
        )
        doc.add_paragraph(
            "Threats flagged by both detection layers received elevated confidence "
            "scoring. Final severity classification incorporated threat category, "
            "affected scope, and temporal persistence."
        )

    def _set_cell_bg(self, cell, hex_color: str):
        shading_elm = cell._element.get_or_add_tcPr()
        shading = shading_elm.makeelement(qn("w:shd"), {
            qn("w:fill"): hex_color,
            qn("w:val"): "clear",
        })
        shading_elm.append(shading)
